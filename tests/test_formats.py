from __future__ import annotations

import json
from pathlib import Path

import openpyxl
import pymupdf
from docx import Document

from aviation_data.extraction import (
    _csv_to_markdown,
    _docx_to_markdown,
    _extraction_quality_flags,
    _html_to_markdown,
    _json_to_markdown,
    _pdf_to_markdown,
    _xlsx_to_markdown,
    _xml_to_markdown,
)
from aviation_data.registry import load_registry

ROOT = Path(__file__).resolve().parents[1]


def test_structured_text_extractors() -> None:
    json_markdown, json_layout = _json_to_markdown(
        json.dumps({"aircraft": "Example", "engines": 2})
    )
    assert "| aircraft | Example |" in json_markdown
    assert json_layout["extractor"] == "json"

    csv_markdown, csv_layout = _csv_to_markdown("code,name\nLTFM,Istanbul\n")
    assert "| LTFM | Istanbul |" in csv_markdown
    assert csv_layout["extractor"] == "csv"

    xml_markdown, xml_layout = _xml_to_markdown(
        "<airport><code>LTFM</code><name>Istanbul</name></airport>"
    )
    assert "| airport/code | LTFM |" in xml_markdown
    assert xml_layout["extractor"] == "xml.etree"


def test_binary_extractors(tmp_path: Path) -> None:
    docx_path = tmp_path / "fixture.docx"
    document = Document()
    document.add_heading("Maintenance", level=1)
    document.add_paragraph("The inspection interval is recorded in the approved programme.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Item"
    table.cell(0, 1).text = "Interval"
    table.cell(1, 0).text = "Inspection"
    table.cell(1, 1).text = "100 hours"
    document.save(docx_path)
    docx_markdown, docx_layout = _docx_to_markdown(docx_path)
    assert "# Maintenance" in docx_markdown
    assert "100 hours" in docx_markdown
    assert docx_layout["extractor"] == "python-docx"

    xlsx_path = tmp_path / "fixture.xlsx"
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "Runways"
    sheet.append(["Designator", "Length"])
    sheet.append(["18", 3000])
    workbook.save(xlsx_path)
    xlsx_markdown, xlsx_layout = _xlsx_to_markdown(xlsx_path)
    assert "# Runways" in xlsx_markdown
    assert "| 18 | 3000 |" in xlsx_markdown
    assert xlsx_layout["extractor"] == "openpyxl"

    pdf_path = tmp_path / "fixture.pdf"
    pdf = pymupdf.open()
    page = pdf.new_page()
    page.insert_text(
        (72, 72),
        "Airport operations use declared distances for runway planning and safe performance.",
    )
    pdf.save(pdf_path)
    pdf.close()
    pdf_markdown, pdf_layout = _pdf_to_markdown(pdf_path)
    assert "Airport operations" in pdf_markdown
    assert pdf_layout["extractor"] == "pymupdf4llm"


def test_mediawiki_extraction_preserves_dom_order_and_prunes_boilerplate() -> None:
    registry = load_registry(ROOT / "configs" / "sources.yaml")
    source = next(
        item for item in registry.sources if item.source_id == "wikipedia_en_aviation_api"
    )
    markup = """
    <div class="mw-parser-output">
      <table class="infobox">
        <tr><th>Engine</th><td>Turbofan</td></tr>
      </table>
      <p>Lead paragraph with an exact fact.<sup class="reference">[1]</sup></p>
      <h2 id="Design">Design</h2>
      <p>The compressor raises the air pressure.</p>
      <dl>
        <dt>Pressure ratio</dt>
        <dd>
          <span class="mwe-math-element mwe-math-element-block">
            <span style="display: none">
              <math alttext="{\\displaystyle p_2/p_1}">
                <annotation>{\\displaystyle p_2/p_1}</annotation>
              </math>
            </span>
            <img alt="{\\displaystyle p_2/p_1}">
          </span>
        </dd>
      </dl>
      <ul>
        <li>Primary item<ul><li>Nested item</li></ul></li>
      </ul>
      <table class="wikitable">
        <tr><th>Stage</th><th>Purpose</th></tr>
        <tr><td>Compressor</td><td>Raise pressure</td></tr>
      </table>
      <h2 id="See_also">See also</h2>
      <ul><li>Unrelated article</li></ul>
      <h2 id="References">References</h2>
      <div class="mw-references-wrap"><p>Reference details</p></div>
      <table class="navbox"><tr><td>Navigation template noise</td></tr></table>
    </div>
    """

    markdown, layout = _html_to_markdown(markup, source, title="Aircraft engine")

    expected_order = (
        "# Aircraft engine",
        "| Engine | Turbofan |",
        "Lead paragraph with an exact fact.",
        "## Design",
        "The compressor raises the air pressure.",
        "**Pressure ratio**",
        "$$p_2/p_1$$",
        "- Primary item",
        "  - Nested item",
        "| Stage | Purpose |",
    )
    positions = [markdown.index(value) for value in expected_order]
    assert positions == sorted(positions)
    assert "[1]" not in markdown
    assert "Unrelated article" not in markdown
    assert "Reference details" not in markdown
    assert "Navigation template noise" not in markdown
    assert "| Stage | Purpose |\n| --- | --- |\n| Compressor | Raise pressure |" in markdown
    assert layout["extraction_profile"] == "mediawiki_article_v1"
    assert layout["excluded_sections"] == ["See also", "References"]
    assert layout["diagnostics"]["maximum_consecutive_headings"] == 1
    assert layout["diagnostics"]["normalized_math_expressions"] == 1
    assert len(layout["tables"]) == 2


def test_mediawiki_extraction_removes_turkish_noncontent_sections() -> None:
    registry = load_registry(ROOT / "configs" / "sources.yaml")
    source = next(
        item for item in registry.sources if item.source_id == "wikipedia_tr_airlines_api"
    )
    markup = """
    <div class="mw-parser-output">
      <p>Havayolu 2020 yılında faaliyete başladı.</p>
      <h2 id="Tarihçe">Tarihçe</h2>
      <p>İlk tarifeli uçuşunu Ankara'dan gerçekleştirdi.</p>
      <h2 id="Ayrıca_bakınız">Ayrıca bakınız</h2>
      <ul><li>Başka havayolları listesi</li></ul>
      <h2 id="Kaynakça">Kaynakça</h2>
      <p>Kaynak ayrıntıları</p>
      <h2 id="Dış_bağlantılar">Dış bağlantılar</h2>
      <p>Şirket sitesi</p>
    </div>
    """

    markdown, layout = _html_to_markdown(markup, source, title="Örnek Havayolu")

    assert markdown.index("Havayolu 2020") < markdown.index("## Tarihçe")
    assert markdown.index("## Tarihçe") < markdown.index("İlk tarifeli")
    assert "Başka havayolları" not in markdown
    assert "Kaynak ayrıntıları" not in markdown
    assert "Şirket sitesi" not in markdown
    assert layout["excluded_sections"] == [
        "Ayrıca bakınız",
        "Kaynakça",
        "Dış bağlantılar",
    ]


def test_mediawiki_quality_flag_targets_front_loaded_heading_failure() -> None:
    registry = load_registry(ROOT / "configs" / "sources.yaml")
    source = next(
        item for item in registry.sources if item.source_id == "wikipedia_en_aviation_api"
    )
    broken_markup = """
    <div class="mw-parser-output">
      <h2>First</h2><h3>Second</h3><h3>Third</h3>
      <p>The first body paragraph appears after a detached heading group.</p>
    </div>
    """
    valid_markup = """
    <div class="mw-parser-output">
      <p>Lead paragraph.</p>
      <h2>Types</h2><h3>Inline</h3><h4>Air-cooled</h4>
      <p>Technical content resumes after legitimate nested headings.</p>
    </div>
    """

    broken_markdown, broken_layout = _html_to_markdown(broken_markup, source, title="Broken")
    valid_markdown, valid_layout = _html_to_markdown(valid_markup, source, title="Valid")

    assert "detached_heading_run" in _extraction_quality_flags(broken_markdown, broken_layout)
    assert "detached_heading_run" not in _extraction_quality_flags(valid_markdown, valid_layout)
