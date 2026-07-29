from __future__ import annotations

import csv
import html
import json
import re
import xml.etree.ElementTree as ET
from datetime import date
from html.parser import HTMLParser
from pathlib import Path
from typing import Any
from urllib.parse import urlparse

from aviation_data.ids import normalize_text, sha256_text, stable_id, tokens
from aviation_data.io import read_jsonl, write_json, write_jsonl, write_parquet_if_available
from aviation_data.models import (
    DocumentRecord,
    Language,
    SourceDefinition,
    SourceRecord,
    SourceRegistry,
)
from aviation_data.registry import source_index


class ExtractionError(RuntimeError):
    pass


GENERIC_HTML_PROFILE = "generic_html_v2"
MEDIAWIKI_ARTICLE_PROFILE = "mediawiki_article_v1"

BASE_HTML_REMOVE_SELECTORS = (
    "script",
    "style",
    "nav",
    "footer",
    "aside",
    "form",
    "noscript",
)

MEDIAWIKI_REMOVE_SELECTORS = (
    ".navbox",
    ".navbar",
    ".vertical-navbox",
    ".authority-control",
    ".mw-references-wrap",
    ".reflist",
    ".metadata",
    ".noprint",
    ".shortdescription",
    ".ambox",
    ".mbox-small",
    ".sistersitebox",
    ".portalbox",
    ".catlinks",
    ".mw-editsection",
    "sup.reference",
    '[role="navigation"]',
)

MEDIAWIKI_EXCLUDED_SECTIONS = {
    "bibliography",
    "citations",
    "external links",
    "further reading",
    "general references",
    "notes",
    "references",
    "see also",
    "sources",
    "ayrıca bakınız",
    "dipnotlar",
    "dış bağlantılar",
    "ileri okuma",
    "kaynaklar",
    "kaynakça",
    "konuyla ilgili yayınlar",
    "notlar",
}


class _ReadableHTML(HTMLParser):
    block_tags = {
        "article",
        "blockquote",
        "br",
        "dd",
        "div",
        "dl",
        "dt",
        "h1",
        "h2",
        "h3",
        "h4",
        "h5",
        "h6",
        "li",
        "main",
        "p",
        "section",
        "td",
        "th",
        "tr",
    }
    ignored = {"script", "style", "nav", "footer", "aside", "noscript", "form"}

    def __init__(self) -> None:
        super().__init__()
        self.parts: list[str] = []
        self.ignored_depth = 0
        self.heading_prefix: str | None = None

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        if tag in self.ignored:
            self.ignored_depth += 1
        if self.ignored_depth:
            return
        if tag in self.block_tags:
            self.parts.append("\n")
        if re.fullmatch(r"h[1-6]", tag):
            self.heading_prefix = "#" * int(tag[1]) + " "
            self.parts.append(self.heading_prefix)
        elif tag == "li":
            self.parts.append("- ")
        elif tag in {"td", "th"}:
            self.parts.append("| ")

    def handle_endtag(self, tag: str) -> None:
        if tag in self.ignored and self.ignored_depth:
            self.ignored_depth -= 1
            return
        if self.ignored_depth:
            return
        if tag in self.block_tags:
            self.parts.append("\n")

    def handle_data(self, data: str) -> None:
        if not self.ignored_depth:
            self.parts.append(html.unescape(data))

    def markdown(self) -> str:
        text = "".join(self.parts)
        text = re.sub(r"[ \t]+", " ", text)
        text = re.sub(r" *\n *", "\n", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text


def _decode(payload: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "windows-1254", "latin-1"):
        try:
            return payload.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise ExtractionError("unsupported text encoding")


def _html_profile(source: SourceDefinition) -> str:
    return source.extraction.profile if source.extraction else GENERIC_HTML_PROFILE


def _normalize_heading(value: str) -> str:
    value = re.sub(r"\s+", " ", value).strip().rstrip(":")
    return value.casefold()


def _nearest_ancestor(node: Any, tag: str, root: Any) -> Any | None:
    current = node.parent
    while current is not None:
        if current.tag == tag:
            return current
        if current.mem_id == root.mem_id:
            break
        current = current.parent
    return None


def _inside(node: Any, tags: set[str], root: Any) -> bool:
    current = node.parent
    while current is not None:
        if current.tag in tags:
            return True
        if current.mem_id == root.mem_id:
            break
        current = current.parent
    return False


def _direct_list_text(node: Any) -> str:
    parts: list[str] = []
    child = node.child
    while child is not None:
        if child.tag not in {"dl", "ul", "ol"}:
            value = child.text(separator=" ", strip=True)
            if value:
                parts.append(value)
        child = child.next
    return re.sub(r"\s+", " ", " ".join(parts)).strip()


def _direct_definition_text(node: Any) -> str:
    parts: list[str] = []
    child = node.child
    while child is not None:
        if child.tag not in {
            "blockquote",
            "dl",
            "ol",
            "p",
            "pre",
            "table",
            "ul",
        }:
            value = child.text(separator=" ", strip=True)
            if value:
                parts.append(value)
        child = child.next
    return re.sub(r"\s+", " ", " ".join(parts)).strip()


def _list_depth(node: Any, root: Any) -> int:
    depth = 0
    current = node.parent
    while current is not None:
        if current.tag == "li":
            depth += 1
        if current.mem_id == root.mem_id:
            break
        current = current.parent
    return depth


def _escape_table_cell(value: str) -> str:
    value = re.sub(r"\s+", " ", value).strip()
    return value.replace("\\", "\\\\").replace("|", "\\|")


def _html_table_rows(table: Any) -> list[list[str]]:
    rows: list[list[str]] = []
    for row in table.css("tr"):
        nearest_table = _nearest_ancestor(row, "table", table)
        if nearest_table is not None and nearest_table.mem_id != table.mem_id:
            continue
        cells = []
        for cell in row.css("th,td"):
            nearest_row = _nearest_ancestor(cell, "tr", row)
            if nearest_row is not None and nearest_row.mem_id != row.mem_id:
                continue
            cells.append(_escape_table_cell(cell.text(separator=" ", strip=True)))
        if cells:
            rows.append(cells)
    if not rows:
        return []
    width = max(len(row) for row in rows)
    return [row + [""] * (width - len(row)) for row in rows]


def _maximum_heading_run(kinds: list[str]) -> int:
    maximum = 0
    current = 0
    for kind in kinds:
        if kind == "heading":
            current += 1
            maximum = max(maximum, current)
        else:
            current = 0
    return maximum


def _mediawiki_math_text(node: Any) -> str:
    math = node.css_first("math")
    value = math.attributes.get("alttext", "") if math is not None else ""
    value = re.sub(r"^\{\\(?:display|text)style\s*", "", value).strip()
    if value.endswith("}"):
        value = value[:-1].rstrip()
    if not value:
        value = node.text(separator=" ", strip=True)
    delimiter = "$$" if "mwe-math-element-block" in node.attributes.get("class", "") else "$"
    return f"{delimiter}{value}{delimiter}"


def _html_to_markdown(
    text: str,
    source: SourceDefinition,
    *,
    title: str | None = None,
) -> tuple[str, dict[str, Any]]:
    profile = _html_profile(source)
    try:
        from selectolax.parser import HTMLParser as SelectolaxParser
    except ImportError as exc:
        if profile == MEDIAWIKI_ARTICLE_PROFILE:
            raise ExtractionError(
                "mediawiki_article_v1 extraction requires the 'formats' extra"
            ) from exc
        parser = _ReadableHTML()
        parser.feed(text)
        return parser.markdown(), {
            "extractor": "stdlib.html.parser",
            "extraction_profile": profile,
        }

    tree = SelectolaxParser(text)
    content_selector = source.selectors.get("content")
    root = tree.css_first(content_selector) if content_selector else None
    root = root or tree.css_first("main") or tree.css_first("article") or tree.body
    if root is None:
        raise ExtractionError("HTML has no readable content root")

    original_text_characters = len(root.text(separator=" ", strip=True))
    normalized_math_expressions = 0
    if profile == MEDIAWIKI_ARTICLE_PROFILE:
        for node in root.css(".mwe-math-element"):
            node.replace_with(_mediawiki_math_text(node))
            normalized_math_expressions += 1

    remove_selectors = list(BASE_HTML_REMOVE_SELECTORS)
    if profile == MEDIAWIKI_ARTICLE_PROFILE:
        remove_selectors.extend(MEDIAWIKI_REMOVE_SELECTORS)
    removed_by_selector: dict[str, dict[str, int]] = {}
    for selector in remove_selectors:
        nodes = root.css(selector)
        if not nodes:
            continue
        removed_elements = 0
        removed_characters = 0
        for node in nodes:
            removed_characters += len(node.text(separator=" ", strip=True))
            removed_elements += 1
            node.decompose()
        removed_by_selector[selector] = {
            "elements": removed_elements,
            "text_characters": removed_characters,
        }

    chunks: list[str] = []
    kinds: list[str] = []
    tables: list[list[list[str]]] = []
    excluded_sections: list[str] = []
    excluded_section_blocks = 0
    excluded_section_characters = 0
    skipping_section = False
    block_counts: dict[str, int] = {}

    if profile == MEDIAWIKI_ARTICLE_PROFILE and title:
        normalized_title = re.sub(r"\s+", " ", title).strip()
        if normalized_title:
            chunks.append(f"# {normalized_title}")
            kinds.append("title")
            block_counts["title"] = 1

    block_tags = {
        "blockquote",
        "dd",
        "dt",
        "figcaption",
        "h1",
        "h2",
        "h3",
        "h4",
        "h5",
        "h6",
        "li",
        "p",
        "pre",
        "table",
    }
    for node in root.traverse():
        if node.tag not in block_tags:
            continue

        if node.tag == "h2":
            heading = re.sub(r"\s+", " ", node.text(separator=" ", strip=True)).strip()
            normalized_heading = _normalize_heading(heading)
            if (
                profile == MEDIAWIKI_ARTICLE_PROFILE
                and normalized_heading in MEDIAWIKI_EXCLUDED_SECTIONS
            ):
                skipping_section = True
                excluded_sections.append(heading)
                excluded_section_blocks += 1
                excluded_section_characters += len(heading)
                continue
            skipping_section = False

        if skipping_section:
            excluded_section_blocks += 1
            excluded_section_characters += len(node.text(separator=" ", strip=True))
            continue

        if node.tag != "table" and _inside(node, {"table"}, root):
            continue
        if node.tag in {"p", "blockquote", "dd", "dt", "figcaption", "pre"} and _inside(
            node, {"li"}, root
        ):
            continue

        if node.tag == "table":
            if _inside(node, {"table"}, root):
                continue
            rows = _html_table_rows(node)
            if rows:
                width = len(rows[0])
                tables.append(rows)
                table_lines = [
                    "| " + " | ".join(rows[0]) + " |",
                    "| " + " | ".join(["---"] * width) + " |",
                    *["| " + " | ".join(row) + " |" for row in rows[1:]],
                ]
                chunks.append("\n".join(table_lines))
                kinds.append("table")
                block_counts["table"] = block_counts.get("table", 0) + 1
        else:
            value = (
                _direct_list_text(node)
                if node.tag == "li"
                else (
                    _direct_definition_text(node)
                    if node.tag in {"dd", "dt"}
                    else re.sub(r"\s+", " ", node.text(separator=" ", strip=True)).strip()
                )
            )
            if not value:
                continue
            if re.fullmatch(r"h[1-6]", node.tag):
                value = f"{'#' * int(node.tag[1])} {value}"
                kind = "heading"
            elif node.tag == "li":
                value = f"{'  ' * _list_depth(node, root)}- {value}"
                kind = "list_item"
            elif node.tag == "dt":
                value = f"**{value}**"
                kind = "definition_term"
            elif node.tag == "dd":
                kind = "paragraph"
            elif node.tag == "blockquote":
                value = f"> {value}"
                kind = "paragraph"
            elif node.tag == "figcaption":
                value = f"Figure: {value}"
                kind = "caption"
            elif node.tag == "pre":
                value = f"```\n{value}\n```"
                kind = "preformatted"
            else:
                kind = "paragraph"
            chunks.append(value)
            kinds.append(kind)
            block_counts[kind] = block_counts.get(kind, 0) + 1

    headings_before_first_paragraph = 0
    headings_before_first_body_block = 0
    for kind in kinds:
        if kind == "paragraph":
            break
        if kind == "heading":
            headings_before_first_paragraph += 1
    for kind in kinds:
        if kind in {
            "caption",
            "definition_term",
            "list_item",
            "paragraph",
            "preformatted",
            "table",
        }:
            break
        if kind == "heading":
            headings_before_first_body_block += 1
    markdown = "\n\n".join(chunks)
    return "\n\n".join(chunks), {
        "extractor": "selectolax",
        "extraction_profile": profile,
        "content_selector": content_selector or "main/article/body",
        "tables": tables,
        "removed_by_selector": removed_by_selector,
        "excluded_sections": excluded_sections,
        "diagnostics": {
            "original_text_characters": original_text_characters,
            "retained_markdown_characters": len(markdown),
            "excluded_section_blocks": excluded_section_blocks,
            "excluded_section_text_characters": excluded_section_characters,
            "block_counts": block_counts,
            "normalized_math_expressions": normalized_math_expressions,
            "headings_before_first_paragraph": headings_before_first_paragraph,
            "headings_before_first_body_block": headings_before_first_body_block,
            "maximum_consecutive_headings": _maximum_heading_run(kinds),
        },
    }


def _extraction_quality_flags(canonical: str, layout: dict[str, Any]) -> list[str]:
    flags: list[str] = []
    token_count = len(tokens(canonical))
    if token_count > 20_000:
        flags.append("oversized_document")

    nonempty_lines = [line.strip() for line in canonical.splitlines() if line.strip()]
    list_lines = [line for line in nonempty_lines if re.match(r"^\s*-\s+", line)]
    if token_count > 1_000 and len(list_lines) / max(1, len(nonempty_lines)) > 0.75:
        flags.append("list_heavy_document")

    substantive_lines = [line for line in nonempty_lines if len(line) >= 60]
    if len(substantive_lines) >= 20:
        duplicate_ratio = 1 - len(set(substantive_lines)) / len(substantive_lines)
        if duplicate_ratio > 0.15:
            flags.append("excessive_duplicate_lines")

    diagnostics = layout.get("diagnostics", {})
    if int(diagnostics.get("headings_before_first_body_block", 0)) > 2:
        flags.append("detached_heading_run")
    if layout.get("extraction_profile") == MEDIAWIKI_ARTICLE_PROFILE and re.search(
        r"\b(?:Authority control databases|Otorite kontrolü)\b",
        canonical,
        re.IGNORECASE,
    ):
        flags.append("html_boilerplate_remaining")
    return flags


def _flatten_json(value: Any, prefix: str = "") -> list[tuple[str, str]]:
    rows: list[tuple[str, str]] = []
    if isinstance(value, dict):
        for key, child in value.items():
            path = f"{prefix}.{key}" if prefix else str(key)
            rows.extend(_flatten_json(child, path))
    elif isinstance(value, list):
        for index, child in enumerate(value):
            rows.extend(_flatten_json(child, f"{prefix}[{index}]"))
    else:
        rows.append((prefix, "" if value is None else str(value)))
    return rows


def _json_to_markdown(text: str) -> tuple[str, dict[str, Any]]:
    value = json.loads(text)
    rows = _flatten_json(value)
    markdown = ["# Structured record", "", "| Field | Value |", "|---|---|"]
    for key, item in rows:
        safe_key = key.replace("|", "\\|")
        safe_item = item.replace("|", "\\|")
        markdown.append(f"| {safe_key} | {safe_item} |")
    return "\n".join(markdown), {"extractor": "json", "value": value, "rows": rows}


def _csv_to_markdown(text: str) -> tuple[str, dict[str, Any]]:
    dialect = csv.Sniffer().sniff(text[:4096])
    rows = list(csv.reader(text.splitlines(), dialect))
    if not rows:
        raise ExtractionError("empty CSV")
    width = max(len(row) for row in rows)
    rows = [row + [""] * (width - len(row)) for row in rows]
    markdown = ["# Table", "", "| " + " | ".join(rows[0]) + " |"]
    markdown.append("| " + " | ".join(["---"] * width) + " |")
    markdown.extend("| " + " | ".join(row) + " |" for row in rows[1:])
    return "\n".join(markdown), {"extractor": "csv", "rows": rows}


def _xml_to_markdown(text: str) -> tuple[str, dict[str, Any]]:
    root = ET.fromstring(text)
    rows: list[tuple[str, str]] = []

    def walk(element: ET.Element, path: str) -> None:
        tag = element.tag.rsplit("}", maxsplit=1)[-1]
        current = f"{path}/{tag}" if path else tag
        value = " ".join(part.strip() for part in element.itertext() if part.strip())
        if value and len(element) == 0:
            rows.append((current, value))
        for child in element:
            walk(child, current)

    walk(root, "")
    markdown = ["# XML record", "", "| Element | Value |", "|---|---|"]
    markdown.extend(f"| {path} | {value} |" for path, value in rows)
    return "\n".join(markdown), {"extractor": "xml.etree", "rows": rows}


def _docx_to_markdown(path: Path) -> tuple[str, dict[str, Any]]:
    try:
        from docx import Document
    except ImportError as exc:
        raise ExtractionError("DOCX extraction requires the 'formats' extra") from exc
    document = Document(path)
    chunks: list[str] = []
    paragraphs = []
    for paragraph in document.paragraphs:
        value = paragraph.text.strip()
        if not value:
            continue
        style = paragraph.style.name if paragraph.style else ""
        heading = re.search(r"Heading\s+([1-6])", style, re.IGNORECASE)
        output = f"{'#' * int(heading.group(1))} {value}" if heading else value
        chunks.append(output)
        paragraphs.append({"style": style, "text": value})
    tables = []
    for table_index, table in enumerate(document.tables, start=1):
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        if not rows:
            continue
        width = max(len(row) for row in rows)
        rows = [row + [""] * (width - len(row)) for row in rows]
        chunks.extend(
            [
                f"## Table {table_index}",
                "| " + " | ".join(rows[0]) + " |",
                "| " + " | ".join(["---"] * width) + " |",
                *["| " + " | ".join(row) + " |" for row in rows[1:]],
            ]
        )
        tables.append(rows)
    return "\n\n".join(chunks), {
        "extractor": "python-docx",
        "paragraphs": paragraphs,
        "tables": tables,
    }


def _xlsx_to_markdown(path: Path) -> tuple[str, dict[str, Any]]:
    try:
        import openpyxl
    except ImportError as exc:
        raise ExtractionError("XLSX extraction requires the 'formats' extra") from exc
    workbook = openpyxl.load_workbook(path, read_only=True, data_only=True)
    chunks: list[str] = []
    sheets: dict[str, list[list[str]]] = {}
    for sheet in workbook.worksheets:
        rows = [
            ["" if cell is None else str(cell) for cell in row]
            for row in sheet.iter_rows(values_only=True)
        ]
        rows = [row for row in rows if any(cell for cell in row)]
        if not rows:
            continue
        width = max(len(row) for row in rows)
        rows = [row + [""] * (width - len(row)) for row in rows]
        chunks.extend(
            [
                f"# {sheet.title}",
                "| " + " | ".join(rows[0]) + " |",
                "| " + " | ".join(["---"] * width) + " |",
                *["| " + " | ".join(row) + " |" for row in rows[1:]],
            ]
        )
        sheets[sheet.title] = rows
    workbook.close()
    if not chunks:
        raise ExtractionError("XLSX contains no non-empty cells")
    return "\n\n".join(chunks), {"extractor": "openpyxl", "sheets": sheets}


def _pdf_to_markdown(path: Path) -> tuple[str, dict[str, Any]]:
    try:
        import pymupdf
        import pymupdf4llm
    except ImportError as exc:
        raise ExtractionError("PDF extraction requires the 'formats' extra") from exc
    chunks = pymupdf4llm.to_markdown(str(path), page_chunks=True, write_images=False)
    if isinstance(chunks, str):
        markdown = chunks
        layout = [{"page": None, "text": chunks}]
    else:
        pages = []
        layout = []
        for index, chunk in enumerate(chunks, start=1):
            value = chunk.get("text", "")
            pages.append(f"<!-- page:{index} -->\n\n{value}")
            layout.append(chunk)
        markdown = "\n\n".join(pages)
    if len(tokens(markdown)) < 10:
        document = pymupdf.open(path)
        ocr_pages = []
        for index, page in enumerate(document, start=1):
            text_page = page.get_textpage_ocr(language="eng+tur", dpi=300, full=True)
            ocr_pages.append(f"<!-- page:{index} -->\n\n{page.get_text(textpage=text_page)}")
        document.close()
        ocr_text = "\n\n".join(ocr_pages)
        if len(tokens(ocr_text)) > len(tokens(markdown)):
            markdown = ocr_text
            layout = [{"extractor": "pymupdf-ocr", "languages": "eng+tur"}]
    return markdown, {"extractor": "pymupdf4llm", "pages": layout}


def _extract_payload(
    path: Path, record: SourceRecord, source: SourceDefinition
) -> tuple[str, dict[str, Any]]:
    mime = record.detected_mime.casefold()
    suffix = Path(urlparse(record.canonical_url).path).suffix.casefold()
    native = record.native_format.casefold()
    if mime == "application/pdf" or suffix == ".pdf" or native == "pdf":
        return _pdf_to_markdown(path)
    if suffix == ".docx" or native == "docx":
        return _docx_to_markdown(path)
    if suffix == ".xlsx" or native == "xlsx":
        return _xlsx_to_markdown(path)
    text = _decode(path.read_bytes())
    if mime == "text/html" or suffix in {".html", ".htm"} or native == "html":
        return _html_to_markdown(text, source, title=record.title)
    if mime in {"application/json", "text/json"} or suffix == ".json" or native == "json":
        return _json_to_markdown(text)
    if mime in {"application/xml", "text/xml"} or suffix == ".xml" or native == "xml":
        return _xml_to_markdown(text)
    if mime in {"text/csv", "application/csv"} or suffix == ".csv" or native == "csv":
        return _csv_to_markdown(text)
    if mime.startswith("text/") or suffix in {".md", ".markdown", ".txt"}:
        return text, {"extractor": "text", "encoding": "utf-8-normalized"}
    raise ExtractionError(f"unsupported MIME/native format: {mime}/{native}")


def _table_rows(layout: dict[str, Any]) -> list[tuple[str, list[dict[str, str]]]]:
    collections: list[tuple[str, list[list[str]]]] = []
    extractor = layout.get("extractor")
    if extractor in {"json", "xml.etree"}:
        return [
            (
                "fields",
                [
                    {"field": str(field), "value": str(value)}
                    for field, value in layout.get("rows", [])
                ],
            )
        ]
    if extractor == "csv":
        collections.append(("table", layout.get("rows", [])))
    for index, rows in enumerate(layout.get("tables", []), start=1):
        collections.append((f"table_{index}", rows))
    for name, rows in layout.get("sheets", {}).items():
        collections.append((f"sheet_{name}", rows))
    output = []
    for name, rows in collections:
        if not rows:
            continue
        width = max(len(row) for row in rows)
        header = [
            (str(value).strip() or f"column_{index + 1}")
            for index, value in enumerate(rows[0] + [""] * (width - len(rows[0])))
        ]
        deduplicated = []
        counts: dict[str, int] = {}
        for value in header:
            counts[value] = counts.get(value, 0) + 1
            deduplicated.append(value if counts[value] == 1 else f"{value}_{counts[value]}")
        output.append(
            (
                name,
                [
                    {
                        key: str(value)
                        for key, value in zip(
                            deduplicated,
                            row + [""] * (width - len(row)),
                            strict=True,
                        )
                    }
                    for row in rows[1:]
                ],
            )
        )
    return output


def _language(text: str, declared: list[Language]) -> Language:
    if len(declared) == 1:
        return declared[0]
    lowered = text.casefold()
    turkish_markers = sum(lowered.count(char) for char in "çğıöşü")
    turkish_words = sum(
        len(re.findall(rf"\b{word}\b", lowered))
        for word in ("bir", "ve", "için", "hava", "olan", "ile", "pist")
    )
    english_words = sum(
        len(re.findall(rf"\b{word}\b", lowered))
        for word in ("the", "and", "for", "aircraft", "runway", "with")
    )
    if turkish_markers + turkish_words > english_words:
        return Language.TURKISH
    if english_words:
        return Language.ENGLISH
    return Language.UNDETERMINED


def _source_as_of(record: SourceRecord) -> date | None:
    revision_timestamp = record.fetch_recipe.get("revision_timestamp")
    if not isinstance(revision_timestamp, str):
        return None
    try:
        return date.fromisoformat(revision_timestamp[:10])
    except ValueError:
        return None


def extract_sources(
    registry: SourceRegistry,
    data_dir: Path,
) -> tuple[list[DocumentRecord], list[dict[str, str]]]:
    records = read_jsonl(data_dir / "manifests" / "source_records.jsonl", SourceRecord)
    definitions = source_index(registry)
    documents: list[DocumentRecord] = []
    errors: list[dict[str, str]] = []
    for record in records:
        source = definitions.get(record.registry_source_id)
        if source is None:
            errors.append(
                {
                    "source_record_id": record.source_record_id,
                    "error": "source missing from current registry",
                }
            )
            continue
        source_path = data_dir / record.storage_path
        try:
            markdown, layout = _extract_payload(source_path, record, source)
            canonical = normalize_text(markdown)
            flags = []
            if len(tokens(canonical)) < 10:
                flags.append("very_short")
            replacement_ratio = canonical.count("\ufffd") / max(1, len(canonical))
            if replacement_ratio > 0.001:
                flags.append("encoding_replacement_noise")
            flags.extend(_extraction_quality_flags(canonical, layout))
            digest = sha256_text(canonical)
            document_id = stable_id(
                "doc", record.registry_source_id, record.source_version, digest, length=32
            )
            artifact_dir = data_dir / "extracted" / "artifacts" / document_id
            canonical_path = artifact_dir / "canonical.md"
            plain_path = artifact_dir / "canonical.txt"
            layout_path = artifact_dir / "layout.json"
            canonical_path.parent.mkdir(parents=True, exist_ok=True)
            canonical_path.write_text(canonical, encoding="utf-8", newline="\n")
            plain = re.sub(r"^#{1,6}\s+", "", canonical, flags=re.MULTILINE)
            plain_path.write_text(plain, encoding="utf-8", newline="\n")
            write_json(layout_path, layout)
            canonical_relative = canonical_path.relative_to(data_dir).as_posix()
            artifact_paths = {
                "source": record.storage_path,
                "canonical_markdown": canonical_relative,
                "canonical_text": plain_path.relative_to(data_dir).as_posix(),
                "layout_json": layout_path.relative_to(data_dir).as_posix(),
            }
            for table_name, rows in _table_rows(layout):
                safe_name = re.sub(r"[^A-Za-z0-9_.-]+", "_", table_name)
                table_path = artifact_dir / f"{safe_name}.parquet"
                if write_parquet_if_available(table_path, rows):
                    artifact_paths[f"table_parquet:{table_name}"] = table_path.relative_to(
                        data_dir
                    ).as_posix()
            documents.append(
                DocumentRecord(
                    document_id=document_id,
                    document_version=record.source_version,
                    variant_group_id=stable_id(
                        "variant",
                        record.registry_source_id,
                        record.canonical_url,
                        length=24,
                    ),
                    title=record.title or Path(urlparse(record.canonical_url).path).stem,
                    language=_language(canonical, record.languages),
                    topics=record.topics,
                    as_of=_source_as_of(record),
                    publisher=record.publisher,
                    source_family=record.source_family,
                    authority_level=record.authority_level,
                    source_record_id=record.source_record_id,
                    source_url=record.canonical_url,
                    native_mime=record.detected_mime,
                    native_format=record.native_format,
                    license_id=record.rights.license_id,
                    attribution=record.rights.attribution,
                    rights_state=record.rights.state,
                    release_derived_text=record.rights.release_derived_text,
                    release_qa=record.rights.release_qa,
                    canonical_path=canonical_relative,
                    canonical_sha256=digest,
                    canonical_char_count=len(canonical),
                    canonical_token_count=len(tokens(canonical)),
                    artifact_paths=artifact_paths,
                    derived_from=[record.source_record_id],
                    quality_flags=flags,
                )
            )
        except Exception as exc:
            errors.append(
                {
                    "source_record_id": record.source_record_id,
                    "source_id": record.registry_source_id,
                    "error": f"{type(exc).__name__}: {exc}",
                }
            )
    documents.sort(key=lambda item: item.document_id)
    write_jsonl(data_dir / "extracted" / "documents.jsonl", documents)
    write_parquet_if_available(data_dir / "extracted" / "documents.parquet", documents)
    write_json(data_dir / "extracted" / "errors.json", errors)
    return documents, errors
